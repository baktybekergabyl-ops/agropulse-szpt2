import shutil
from copy import copy
from datetime import timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from .models import Snapshot
from .parser import key, product_key
from .report import PRODUCT_ALIASES, _national_price

MONTHS_NOMINATIVE = [
    "", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
]
MONTHS_GENITIVE = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]
DISPLAY_BY_SOURCE = {
    product_key(source): display.lower()
    for display, source in PRODUCT_ALIASES.items()
}
AGGREGATE_KEY = product_key("Социально-значимые продовольственные товары")


def _copy_cell_style(source, target) -> None:
    target._style = copy(source._style)
    if source.has_style:
        target.font = copy(source.font)
        target.fill = copy(source.fill)
        target.border = copy(source.border)
        target.alignment = copy(source.alignment)
        target.number_format = source.number_format
        target.protection = copy(source.protection)


def _template_product_key(value: object) -> str:
    name = str(value or "").strip()
    source_name = PRODUCT_ALIASES.get(key(name), name)
    return product_key(source_name)


def create_comparison_workbook(
    template_path: Path,
    current: Snapshot,
    output_dir: Path,
) -> Path:
    if not template_path.exists():
        raise FileNotFoundError(f"Шаблон динамики не найден: {template_path}")
    output_dir.mkdir(parents=True, exist_ok=True)
    date = current.source.observation_date
    output = output_dir / f"Сравнение цен по периодам на {date:%d.%m.%Y}.xlsx"
    shutil.copy2(template_path, output)
    workbook = load_workbook(output)
    sheet_name = str(date.year)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"В шаблоне динамики нет листа {sheet_name}.")
    sheet = workbook[sheet_name]
    date_token = date.strftime("%d.%m")
    date_column = next(
        (
            col for col in range(4, sheet.max_column + 1)
            if date_token in str(sheet.cell(4, col).value or "")
        ),
        None,
    )
    deviation_column = next(
        col for col in range(4, sheet.max_column + 1)
        if str(sheet.cell(3, col).value or "").startswith("Отклонение за год")
    )
    if deviation_column is None:
        raise ValueError("В шаблоне динамики не найдены колонки отклонений.")

    if date_column is None:
        insert_at = deviation_column
        year_merge = next(
            (
                merged for merged in list(sheet.merged_cells.ranges)
                if merged.min_row == 3
                and merged.max_row == 3
                and merged.max_col == insert_at - 1
            ),
            None,
        )
        for address in (
            f"{sheet.cell(3, insert_at).coordinate}:{sheet.cell(4, insert_at).coordinate}",
            f"{sheet.cell(3, insert_at + 1).coordinate}:{sheet.cell(4, insert_at + 1).coordinate}",
        ):
            if address in {str(item) for item in sheet.merged_cells.ranges}:
                sheet.unmerge_cells(address)
        if year_merge:
            year_start = year_merge.min_col
            sheet.unmerge_cells(str(year_merge))
        else:
            year_start = 4
        sheet.insert_cols(insert_at, 1)
        for row in range(1, sheet.max_row + 1):
            _copy_cell_style(sheet.cell(row, insert_at - 1), sheet.cell(row, insert_at))
        source_letter = get_column_letter(insert_at - 1)
        target_letter = get_column_letter(insert_at)
        sheet.column_dimensions[target_letter].width = (
            sheet.column_dimensions[source_letter].width
        )
        sheet.merge_cells(
            start_row=3, start_column=year_start, end_row=3, end_column=insert_at
        )
        sheet.merge_cells(
            start_row=3, start_column=insert_at + 1, end_row=4, end_column=insert_at + 1
        )
        sheet.merge_cells(
            start_row=3, start_column=insert_at + 2, end_row=4, end_column=insert_at + 2
        )
        date_column = insert_at
        deviation_column = insert_at + 1

    sheet.cell(4, date_column).value = (
        f"{MONTHS_NOMINATIVE[date.month]} ({date:%d.%m})"
    )
    for row in range(5, 36):
        item_key = _template_product_key(sheet.cell(row, 2).value)
        sheet.cell(row, date_column).value = _national_price(current, item_key)
        sheet.cell(row, deviation_column).value = current.annual_change.get(item_key)
        sheet.cell(row, deviation_column + 1).value = current.year_change.get(item_key)
    workbook.save(output)
    return output


def _percent(value: float) -> str:
    text = f"{value:.1f}".rstrip("0").rstrip(".").replace(".", ",")
    return f"{text}%"


def _movement(value: float) -> str:
    if value > 0:
        return f"вырос на {_percent(value)}"
    if value < 0:
        return f"снизился на {_percent(abs(value))}"
    return "не изменился"


def _positive_items(values: dict[str, float]) -> list[tuple[str, float]]:
    return sorted(
        [
            (DISPLAY_BY_SOURCE.get(item_key, item_key), value)
            for item_key, value in values.items()
            if item_key != AGGREGATE_KEY and value > 0
        ],
        key=lambda item: (-item[1], item[0]),
    )


def _items_sentence(items: list[tuple[str, float]]) -> str:
    if not items:
        return "Товаров с ростом цен не зафиксировано."
    return "Это " + ", ".join(
        f"{name} - {_percent(value)}" for name, value in items
    ) + "."


def _replace_paragraph(paragraph, text: str, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.bold = bold
    run.font.size = Pt(16)


def create_brief(
    template_path: Path,
    current: Snapshot,
    previous: Snapshot | None,
    output_dir: Path,
) -> Path:
    if not template_path.exists():
        raise FileNotFoundError(f"Шаблон справки не найден: {template_path}")
    output_dir.mkdir(parents=True, exist_ok=True)
    date = current.source.observation_date
    output = output_dir / f"Справка по росту цен СЗПТ на {date:%d.%m.%Y}.docx"
    shutil.copy2(template_path, output)
    document = Document(output)
    previous_date = (
        previous.source.observation_date if previous else date - timedelta(days=7)
    )
    weekly_total = current.week_change.get(AGGREGATE_KEY, 0.0)
    yearly_total = current.year_change.get(AGGREGATE_KEY, 0.0)
    weekly_items = _positive_items(current.week_change)
    yearly_items = _positive_items(current.year_change)

    _replace_paragraph(
        document.paragraphs[2],
        f"(на {date.day} {MONTHS_GENITIVE[date.month]} {date.year} года)",
    )
    _replace_paragraph(
        document.paragraphs[4],
        f"По итогам прошедшей недели ({previous_date:%d.%m.%Y} - "
        f"{date:%d.%m.%Y}) индекс цен {_movement(weekly_total)}.",
    )
    _replace_paragraph(
        document.paragraphs[5],
        f"При этом, в разрезе товаров рост цен отмечен по "
        f"{len(weekly_items)} наименованиям СЗПТ.",
    )
    _replace_paragraph(document.paragraphs[6], _items_sentence(weekly_items))
    _replace_paragraph(
        document.paragraphs[8],
        f"С начала года индекс цен по республике {_movement(yearly_total)}.",
    )
    _replace_paragraph(
        document.paragraphs[9],
        f"В разрезе товаров рост цен отмечен по "
        f"{len(yearly_items)} наименованиям СЗПТ.",
    )
    _replace_paragraph(document.paragraphs[10], _items_sentence(yearly_items))
    document.save(output)
    return output
